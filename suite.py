import time
import sys
import os
from threading import Thread
import configparser
import paramiko
from docx import Document
'''
This is the Machine Class
'''
class Machine:


	#Comments have to be indented correctly?!?
	'''
	This is the constructor (called something else in python?) it takes as
	input and ip address and sets/creates an instance variable for the ip along
	with the creation of a hostname instance variable and lists for rawConfig,
	cleanedConfig, and ip_route.
	'''
	def __init__(self, ip, username, password):
		self.ip = ip
		self.hostname = "S1"
		self.username = username
		self.password = password
		self.rawConfig = []
		self.cleanedConfig = []
		self.ip_route = []


	'''
	This function is used for gathering the config files of the
	machines and the routing table information. The function takes as
	input an object that it is called upon...
	'''
	def gatherConfigs(self):

		ssh = paramiko.SSHClient()
		#In case the machine is not a known host. What if the fingerprint
		#changed though?
		ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
		try:
			ssh.connect(self.ip, port=22, username=self.username, password=self.password, allow_agent=False, look_for_keys=False)
			stdin, stdout, stderr = ssh.exec_command('show run')
			self.rawConfig = stdout.readlines()
			#stdin, stdout, stderr = ssh.exec_command('ip route')
			#self.ip_route = stdout.readlines()
			ssh.close()
			#print(self.rawConfig)
		except:
			print(sys.exc_info()[0])

		self.cleanConfigss()


	def cleanConfigss(self):
		count = 0
		last = ""

		for line in self.rawConfig:
			if count > 2:
				if self.rawConfig[count] != last: #Might be expensive. There are other ways...
					self.cleanedConfig.append(self.rawConfig[count])
					last = self.rawConfig[count]
			count += 1
		self.cleanedConfig[-1] = 'end' # NOT GOOD
		print(self.cleanedConfig[-1])
		print(self.cleanedConfig)
		self.write_cleaned()

	def write_cleaned(self):
		f = open('config' + str(time.time()),'w')
		for i in self.cleanedConfig:
			f.write(i)
		#Should I seperate these?

#-------------------------------------------
def populateMachineList(list_name, machines, username, password):
	ip_list = open(list_name, 'r')
	#Create and put the machine objects in a list.
	for ip in ip_list:
		machines.append(Machine(ip, username, password))

def startThreads(machines, threads):
	for i in machines:
		t1 = Thread(target=i.gatherConfigs, args=())
		t1.start()
		threads.append(t1)
	for t1 in threads:
		t1.join()

def write_lab(machines):
	document = Document()
	for i in machines:
		document.add_heading(i.hostname, 1)
		p = document.add_paragraph()
		for line in i.cleanedConfig:
			p.add_run(line.rstrip('\n'))
		#document.add_paragraph(i.cleanedConfig.rstrip())

	document.save('Lab_Write_Up.docx')

'''
Thanks for the code Chris
'''
def get_config_info():

    # get the path to the config file from the first parameter
    config_file_path = ''

    if len(sys.argv) > 1:
        config_file_path = sys.argv[1]
    else:
        print('Usage: %s /path/to/config.ini' % (sys.argv[0],))
        sys.exit()

    config = configparser.ConfigParser()

    # parse the content of the config file
    with open(config_file_path, 'r') as f:
        config.read_file(f)

    username = config['Lab_Suite']['username']
    password = config['Lab_Suite']['password']

    return username, password


def Main():
    username = ''
    password = ''
    try:
        username, password = get_config_info()
    finally:
        print("Make sure you have the proper config information")
        
	machines = []
	threads = []
	populateMachineList('ips.txt', machines, username, password)
	os.mkdir('Cleaned_Configs')
	os.chdir('Cleaned_Configs')
	startThreads(machines, threads)
	os.chdir('..')
	#Consider having a method here for writting the files to disk.

	write_lab(machines)


#Necessary?
if __name__ == '__main__':
	timez = time.time()
	Main()
	print(time.time() - timez)
